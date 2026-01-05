# ABOUTME: Pytest fixtures for word-ocr tests
# ABOUTME: Provides sample Word documents with embedded images

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Inches
from PIL import Image
import io


@pytest.fixture
def sample_docx_with_images(tmp_path):
    """Create a Word document with embedded images."""
    doc = Document()
    doc.add_heading('Test Document', 0)

    # Create and add first image
    img1 = Image.new('RGB', (100, 100), color='red')
    img1_bytes = io.BytesIO()
    img1.save(img1_bytes, format='PNG')
    img1_bytes.seek(0)

    doc.add_paragraph('First paragraph with image:')
    doc.add_picture(img1_bytes, width=Inches(1.0))

    # Create and add second image
    img2 = Image.new('RGB', (100, 100), color='blue')
    img2_bytes = io.BytesIO()
    img2.save(img2_bytes, format='PNG')
    img2_bytes.seek(0)

    doc.add_paragraph('Second paragraph with image:')
    doc.add_picture(img2_bytes, width=Inches(1.0))

    # Save document
    doc_path = tmp_path / "test_with_images.docx"
    doc.save(doc_path)

    return doc_path


@pytest.fixture
def sample_docx_no_images(tmp_path):
    """Create a Word document without images."""
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('Just text, no images.')

    doc_path = tmp_path / "test_no_images.docx"
    doc.save(doc_path)

    return doc_path


@pytest.fixture
def empty_docx(tmp_path):
    """Create an empty Word document."""
    doc = Document()

    doc_path = tmp_path / "empty.docx"
    doc.save(doc_path)

    return doc_path
